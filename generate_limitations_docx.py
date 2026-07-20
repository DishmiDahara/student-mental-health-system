import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = docx.Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style definitions
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0x11, 0x18, 0x27) # slate-900

    # Title Heading
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("Limitations of the Existing Manual Student Mental Health System")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Academic Blue (#1E3A8A)

    # Introduction
    p_intro = doc.add_paragraph(
        "The traditional manual system for managing student mental health support at the university "
        "faces several critical limitations and operational inefficiencies. These bottlenecks directly "
        "hinder the effectiveness, safety, and reach of the support program, as summarized below:"
    )
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.paragraph_format.space_after = Pt(12)

    # Bulleted Point Form Limitations
    def add_bullet_point(doc, title, detail):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        # Add bold title
        r_title = p.add_run(title + ": ")
        r_title.bold = True
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(12)
        
        # Add detail description
        r_detail = p.add_run(detail)
        r_detail.font.name = 'Times New Roman'
        r_detail.font.size = Pt(12)

    add_bullet_point(doc, "Limited Accessibility", "Operating hours are strictly restricted to office timings (9:00 AM to 4:00 PM, Monday-Friday), leaving students without support during late-night or weekend crises.")
    add_bullet_point(doc, "Significant Response Delays", "Matching students with available counselors manually takes 2 to 5 business days, creating a critical wait time when immediate attention is needed.")
    add_bullet_point(doc, "Lack of Confidentiality and High Stigma", "Mandatory physical office walk-ins and identifiable university email communications expose students to social stigma, deterring many from seeking help.")
    add_bullet_point(doc, "Insecure Physical Record Keeping", "Intake logs and clinical notes are recorded on paper forms stored in physical cabinets, rendering them vulnerable to loss, damage, and slow retrieval.")
    add_bullet_point(doc, "No Automated Crisis Triage", "All counseling requests are handled chronologically without severity-based triage, meaning high-risk students (e.g., severe depression or self-harm) are not prioritized.")
    add_bullet_point(doc, "Absence of Real-Time Interaction", "No immediate coping tools, emotional tracking widgets, or AI chatbot services exist to help students who require instant guidance.")
    add_bullet_point(doc, "Lack of Centralized Analytics", "Administrators cannot evaluate overall mental health trends or counselor workload statistics because data is scattered across paper records and individual spreadsheets.")

    # Save Document
    doc.save("Existing_System_Limitations.docx")
    print("Limitations document created successfully.")

if __name__ == "__main__":
    create_document()
