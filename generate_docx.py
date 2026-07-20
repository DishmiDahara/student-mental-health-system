import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style definitions
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0x11, 0x18, 0x27) # slate-900

    # 1.4 Heading
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("1.4 Existing System Analysis")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Academic Blue (#1E3A8A)

    # Paragraph 1
    p1 = doc.add_paragraph(
        "The current process of managing student mental health inquiries and session schedules within the "
        "university administration is predominantly paper-based and manual. When a student seeks counseling "
        "services, they must initiate a request through formal walk-ins, phone calls, or direct emails to the "
        "counseling center office. This request is received by administrative staff who manually check counselor "
        "calendars and shift timetables to verify availability. Once a slot is coordinated, details are recorded in "
        "physical ledgers or spreadsheet files, and confirmation is sent back to the student."
    )
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(8)

    # Paragraph 2
    p2 = doc.add_paragraph(
        "During counseling sessions, counselors write down notes and assessments on physical intake forms. "
        "These records are then stored in locking filing cabinets in the department. Any follow-up appointments, "
        "escalations, or feedback questionnaires require manual triggers, which frequently lead to long delays, "
        "loss of documentation, and administrative overhead."
    )
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(14)

    # Subheading 1.4.1
    subheading1 = doc.add_paragraph()
    subheading1.paragraph_format.space_before = Pt(14)
    subheading1.paragraph_format.space_after = Pt(8)
    subheading1.paragraph_format.keep_with_next = True
    run_sub1 = subheading1.add_run("1.4.1 Vulnerabilities and Bottlenecks")
    run_sub1.font.name = 'Times New Roman'
    run_sub1.font.size = Pt(13)
    run_sub1.font.bold = True
    run_sub1.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    # List of Bottlenecks
    def add_bullet_item(doc, title, desc):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r_title = p.add_run(title + ": ")
        r_title.bold = True
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(12)
        r_desc = p.add_run(desc)
        r_desc.font.name = 'Times New Roman'
        r_desc.font.size = Pt(12)

    add_bullet_item(doc, "Limited Accessibility", "Counseling services are constrained strictly to standard office hours (e.g., 9:00 AM to 4:00 PM). Students experiencing emotional stress or anxiety during late-night hours have no immediate means of connection or guidance.")
    add_bullet_item(doc, "Delayed Response Time", "Because calendar coordination and counselor assignments are performed manually, the typical turnaround time between a request and an appointment ranges from two to five business days, which is critical in mental health situations.")
    add_bullet_item(doc, "Lack of Privacy & Anonymity", "Physically walking into the counseling center or sending emails using student accounts exposes students to social stigma. This lack of confidentiality deters a large demographic of students from seeking help.")
    add_bullet_item(doc, "Manual Record Management", "Reliance on physical paper files and localized spreadsheets creates a risk of records being misplaced or lost. Furthermore, analyzing long-term student wellness metrics or generating compliance reports is highly inefficient.")
    add_bullet_item(doc, "No Real-Time Support or Triage", "The current system treats all requests chronologically without an automated diagnostic system. This prevents administrators from immediately escalating critical or high-risk cases.")

    # Subheading 1.4.2
    subheading2 = doc.add_paragraph()
    subheading2.paragraph_format.space_before = Pt(16)
    subheading2.paragraph_format.space_after = Pt(10)
    subheading2.paragraph_format.keep_with_next = True
    run_sub2 = subheading2.add_run("1.4.2 Comparative Evaluation Matrix")
    run_sub2.font.name = 'Times New Roman'
    run_sub2.font.size = Pt(13)
    run_sub2.font.bold = True
    run_sub2.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    # Paragraph introducing the table
    p_tbl = doc.add_paragraph(
        "To address these systemic problems, the proposed digital system (MindSpace) introduces "
        "a three-tier automated structure. The table below outlines a comparison between the "
        "existing manual system and the proposed system across critical parameters:"
    )
    p_tbl.paragraph_format.line_spacing = 1.15
    p_tbl.paragraph_format.space_after = Pt(12)

    # Table creation: 3 columns, 7 rows (1 header + 6 content)
    table = doc.add_table(rows=7, cols=3)
    table.autofit = False
    
    # Column widths
    col_widths = [Inches(1.8), Inches(2.35), Inches(2.35)]
    
    # Headers
    headers = ["Feature / Dimension", "Existing Manual System", "Proposed MindSpace System"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "1E3A8A") # Academic dark blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)

    # Table data
    data = [
        ("Availability", 
         "Restricted to physical office hours (Mon-Fri, 9am - 4pm).", 
         "24/7 instantaneous availability via AI Chatbot and web portal."),
        
        ("Privacy & Anonymity", 
         "No anonymity; requires physical visits or academic email tracking, raising stigma concerns.", 
         "Highly anonymous; secure JWT logins, encrypted logs, and peer-to-peer anonymous chats."),
        
        ("Booking & Scheduling", 
         "Manual booking via emails/phone, taking 2 to 5 days to resolve counselor matchings.", 
         "Automated real-time booking showing counselor slots instantly with email confirmation."),
        
        ("Data Security", 
         "Stored in physical folders and locking cabinets; high vulnerability to loss or physical damage.", 
         "Secure cloud database (MongoDB) with role-based access controls and daily backups."),
        
        ("Crisis Triage & Support", 
         "No real-time severity check; cases processed chronologically regardless of stress severity.", 
         "Automated Sentiment Analysis engine on chats to identify and flag high-risk cases for crisis escalation."),
        
        ("Analytical Insights", 
         "Manual review of files to extract statistical data, resulting in poor administration tracking.", 
         "Admin dashboard with wellness tracking logs and student statistics charts to audit support trends.")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row = table.rows[row_idx]
        # Alternating row background colors for readability
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            
            # Align first column left (bold), others left (regular)
            if col_idx == 0:
                run = p.add_run(text)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run = p.add_run(text)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

    # Set column widths for all cells
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Add spacing after the table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Save Document
    doc.save("Section_1_4_Existing_System_Analysis.docx")
    print("Document created successfully.")

if __name__ == "__main__":
    create_document()
